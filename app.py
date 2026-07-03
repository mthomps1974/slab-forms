from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document
from docx.shared import Pt, RGBColor, Twips, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, io, base64, tempfile
from datetime import datetime
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

app = Flask(__name__)
CORS(app)

GMAIL_ADDRESS = os.environ.get('smtp_USER', '')
GMAIL_APP_PASSWORD = os.environ.get('smtp_pass', '')
FIRM_EMAIL = os.environ.get('FIRM_EMAIL', 'info@tflaw.co.uk')
SOL_SIG_PATH = os.path.join(os.path.dirname(__file__), 'sol_sig.png')

def v(d, key):
    return str(d.get(key, '') or '').strip()

def tick(val):
    return 'X' if val else ''

def shade_row(row, hex_col='1F3864'):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_col)
        tcPr.append(shd)

def shade_para(para, hex_col='1F3864'):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_col)
    pPr.append(shd)

def set_col_width(table, col_index, width_cm):
    for row in table.rows:
        row.cells[col_index].width = Cm(width_cm)

def make_client_sig_image(name, sig_data_url=None):
    """Use canvas signature if provided, otherwise render name as italic text."""
    tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
    if sig_data_url and sig_data_url.startswith('data:image'):
        # Decode the canvas PNG from the browser
        header, b64data = sig_data_url.split(',', 1)
        img_bytes = base64.b64decode(b64data)
        with open(tmp.name, 'wb') as f:
            f.write(img_bytes)
    else:
        # Fallback: render name as italic text
        fig, ax = plt.subplots(figsize=(4, 0.9))
        fig.patch.set_facecolor('white')
        ax.set_facecolor('white')
        ax.axis('off')
        ax.text(0.02, 0.5, name,
            fontfamily='DejaVu Serif', style='italic',
            fontsize=28, color='#0d1b2a',
            transform=ax.transAxes, va='center')
        plt.tight_layout(pad=0)
        plt.savefig(tmp.name, format='png', dpi=150, bbox_inches='tight',
            facecolor='white', edgecolor='none')
        plt.close()
    return tmp.name

def section_hdr(doc, text, grey=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    shade_para(p, 'C0C0C0' if grey else '404040')
    run = p.add_run('  ' + text)
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if not grey else RGBColor(0x00, 0x00, 0x00)

def label_value_row(doc, label, value, label_w=6, value_w=10.5):
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.columns[0].width = Cm(label_w)
    table.columns[1].width = Cm(value_w)
    lc = table.rows[0].cells[0]
    vc = table.rows[0].cells[1]
    lc.width = Cm(label_w)
    vc.width = Cm(value_w)
    lr = lc.paragraphs[0].add_run(label)
    lr.font.size = Pt(9)
    vr = vc.paragraphs[0].add_run(value)
    vr.font.size = Pt(9)
    for cell in [lc, vc]:
        cell.paragraphs[0].paragraph_format.space_before = Pt(2)
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)

def two_col_row(doc, l1, v1, l2, v2, lw=4, vw=5):
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    widths = [Cm(lw), Cm(vw), Cm(lw), Cm(vw)]
    cells = table.rows[0].cells
    for i, (cell, w) in enumerate(zip(cells, widths)):
        cell.width = w
        cell.paragraphs[0].paragraph_format.space_before = Pt(2)
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
    cells[0].paragraphs[0].add_run(l1).font.size = Pt(9)
    cells[1].paragraphs[0].add_run(v1).font.size = Pt(9)
    if l2:
        cells[2].paragraphs[0].add_run(l2).font.size = Pt(9)
    cells[3].paragraphs[0].add_run(v2).font.size = Pt(9)

def checkbox_row(doc, items, cols=1):
    if cols == 1:
        for label, checked in items:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Cm(0.5)
            r = p.add_run(('[X]  ' if checked else '[   ]  ') + label)
            r.font.size = Pt(9)
    else:
        half = (len(items) + 1) // 2
        table = doc.add_table(rows=half, cols=2)
        table.style = 'Table Grid'
        for i in range(half):
            for j in range(2):
                idx = i + j * half
                if idx < len(items):
                    label, checked = items[idx]
                    r = table.rows[i].cells[j].paragraphs[0].add_run(
                        ('[X]  ' if checked else '[   ]  ') + label)
                    r.font.size = Pt(9)

def sig_block(doc, sig_label, sig_img_path, date_str, width_inches=3.5):
    doc.add_paragraph()
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    # Header row
    shade_row(table.rows[0], 'E8E8E8')
    hl = table.rows[0].cells[0].paragraphs[0].add_run(sig_label)
    hl.bold = True
    hl.font.size = Pt(9)
    hr = table.rows[0].cells[1].paragraphs[0].add_run('Date:')
    hr.bold = True
    hr.font.size = Pt(9)
    # Signature row
    sig_cell = table.rows[1].cells[0]
    sig_para = sig_cell.paragraphs[0]
    try:
        run = sig_para.add_run()
        run.add_picture(sig_img_path, width=Inches(width_inches))
    except Exception:
        sig_para.add_run('').font.size = Pt(24)
    # Date row
    table.rows[1].cells[1].paragraphs[0].add_run(date_str).font.size = Pt(9)
    # Spacer row
    table.rows[2].cells[0].paragraphs[0].add_run(' ')
    table.rows[2].cells[1].paragraphs[0].add_run(' ')
    doc.add_paragraph()

def decl_items(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.8)
        r = p.add_run('> ' + item)
        r.font.size = Pt(8.5)

def page_break(doc):
    doc.add_page_break()

def make_aa(d, client_sig_path):
    doc = Document()
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(9)

    # ── PAGE 1 ──────────────────────────────────────────────────
    # Header
    hdr_table = doc.add_table(rows=1, cols=3)
    hdr_table.style = 'Table Grid'
    hdr_table.rows[0].cells[0].paragraphs[0].add_run(
        'Civil advice and assistance\nand civil ABWOR Legal Aid Online\nDeclaration').font.size = Pt(11)
    sm_cell = hdr_table.rows[0].cells[1]
    sm_cell.paragraphs[0].add_run('Subject matter:').font.size = Pt(8)
    sm_cell.add_paragraph().add_run(v(d,'subject')).font.size = Pt(9)
    hdr_table.rows[0].cells[2].paragraphs[0].add_run('June 2022\nAA/LAO/CIV').font.size = Pt(8)

    doc.add_paragraph()
    section_hdr(doc, 'A.  Applicant details')

    two_col_row(doc, 'Forename:', v(d,'fname'), 'Surname:', v(d,'lname'))
    label_value_row(doc, 'Date of birth (dd/mm/yyyy):', v(d,'dob'), 5, 11.5)
    two_col_row(doc, 'Contact telephone number:', v(d,'tel'), 'Contact email address:', v(d,'email'))
    two_col_row(doc, 'Contact by email?', v(d,'cemail'), 'National Insurance number:', v(d,'ni'))
    label_value_row(doc, 'If no NI number, reason:', v(d,'ni_reason'))
    two_col_row(doc, 'Home address:', v(d,'haddr')+'\n'+v(d,'hpc'), 'Correspondence address:', v(d,'caddr')+'\n'+v(d,'cpc'), lw=4, vw=5)

    # Comm support
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    pr = p.add_run('Does your client have any communication support needs?')
    pr.bold = True
    pr.font.size = Pt(9)
    comm_val = v(d, 'comm')
    comm_items = [
        ('none/no support needed', 'None' in comm_val or 'none' in comm_val.lower()),
        ('audio', 'Audio' in comm_val),
        ('spoken language translation or interpreter', 'Spoken' in comm_val),
        ('large print', 'Large print' in comm_val),
        ('British Sign Language (BSL)', 'BSL' in comm_val or 'British Sign' in comm_val),
        ('braille', 'Braille' in comm_val or 'braille' in comm_val),
        ('Other', 'Other' in comm_val),
    ]
    checkbox_row(doc, comm_items, cols=2)
    if v(d,'comm_detail'):
        label_value_row(doc, 'Details:', v(d,'comm_detail'))

    page_break(doc)

    # ── PAGE 2 ──────────────────────────────────────────────────
    prevsol = v(d,'prevsol') == 'Yes'
    two_col_row(doc, 'Has your client previously received advice on this matter from any other solicitor?',
                'Yes [X]' if prevsol else 'Yes [  ]',
                '', 'No [X]' if not prevsol else 'No [  ]', lw=11, vw=1.5)
    if prevsol:
        label_value_row(doc, 'If yes, please give details:', v(d,'prevsol_det'))

    section_hdr(doc, 'Applicant/client equality information', grey=True)

    sex = v(d,'sex')
    p = doc.add_paragraph()
    p.add_run('What is your client\'s sex?').bold = True
    checkbox_row(doc, [
        ('Female', sex == 'Female'),
        ('Male', sex == 'Male'),
        ('Prefer not to say', sex == 'Prefer not to say'),
    ])

    dis = v(d,'disabilities')
    p2 = doc.add_paragraph()
    p2.add_run('Disabilities').bold = True
    p2.add_run(' - Does your client have any of the following which have lasted, or are expected to last, at least 12 months?').font.size = Pt(8.5)
    dis_items = [
        ('Deafness or partial hearing loss', 'Deafness' in dis),
        ('Blindness or partial sight loss', 'Blindness' in dis),
        ('Full or partial loss of voice or difficulty speaking', 'voice' in dis.lower()),
        ('Learning disability', 'Learning disability' in dis),
        ('Learning difficulty', 'Learning difficulty' in dis),
        ('Development disorder', 'Development' in dis),
        ('Physical disability', 'Physical' in dis),
        ('Mental health condition', 'Mental' in dis),
        ('Long-term illness, disease or condition', 'Long-term' in dis),
        ('Other condition: please write in below', 'Other condition' in dis),
        ('No known conditions', 'No known' in dis),
        ('Prefer not to say', dis == 'Prefer not to say'),
    ]
    checkbox_row(doc, dis_items)

    page_break(doc)

    # ── PAGE 3 ──────────────────────────────────────────────────
    natid = v(d,'natid')
    p3 = doc.add_paragraph()
    p3.add_run('National identity').bold = True
    nat_items = [
        ('Scottish', 'Scottish' in natid),
        ('English', 'English' in natid),
        ('Northern Irish', 'Northern Irish' in natid),
        ('Welsh', 'Welsh' in natid),
        ('British', 'British' in natid),
        ('Prefer not to say', natid == 'Prefer not to say'),
    ]
    checkbox_row(doc, nat_items, cols=2)

    ethnic = v(d,'ethnic')
    p4 = doc.add_paragraph()
    p4.add_run('Ethnic group').bold = True
    eth_items = [
        ('White Scottish', 'White Scottish' in ethnic),
        ('White Other British', 'White Other British' in ethnic),
        ('White Irish', 'White Irish' in ethnic),
        ('White Polish', 'White Polish' in ethnic),
        ('White Gypsy/Traveller', 'Gypsy' in ethnic),
        ('White Roma', 'Roma' in ethnic),
        ('White Showman/Showwoman', 'Showman' in ethnic),
        ('Pakistani, Scottish Pakistani or British Pakistani', 'Pakistani' in ethnic),
        ('Indian, Scottish Indian or British Indian', 'Indian' in ethnic),
        ('Bangladeshi', 'Bangladeshi' in ethnic),
        ('Chinese, Scottish Chinese or British Chinese', 'Chinese' in ethnic),
        ('African, Scottish African or British African', 'African' in ethnic),
        ('Caribbean or Black', 'Caribbean' in ethnic),
        ('Mixed or multiple ethnic groups', 'Mixed' in ethnic),
        ('Arab, Scottish Arab or British Arab', 'Arab' in ethnic),
        ('Prefer not to say', ethnic == 'Prefer not to say'),
    ]
    checkbox_row(doc, eth_items, cols=2)

    care = v(d,'care')
    p5 = doc.add_paragraph()
    p5.add_run('Care Experience').bold = True
    checkbox_row(doc, [
        ("Currently 'looked after' by Local Authority", 'Currently' in care),
        ("Have previously been 'looked after' by Local Authority", 'previously' in care.lower()),
        ("Never been 'looked after' by Local Authority", 'Never' in care),
        ('Prefer not to say', care == 'Prefer not to say'),
    ])

    page_break(doc)

    # ── PAGE 4 ──────────────────────────────────────────────────
    section_hdr(doc, 'B.  Applicant assistance')
    oa = v(d,'otherassist') == 'Yes'
    two_col_row(doc, 'Do your client have access to any other assistance that provides help with legal costs?',
                'Yes [X]' if oa else 'Yes [  ]', '', 'No [X]' if not oa else 'No [  ]', lw=11, vw=1.5)
    if oa and v(d,'oa_why'):
        label_value_row(doc, 'If this assistance cannot be used, tell us why:', v(d,'oa_why'))

    section_hdr(doc, 'C.  Financial details')
    has_partner = v(d,'partner') == 'Yes'
    two_col_row(doc, 'Does your client live with a spouse/partner?',
                'Yes [X]' if has_partner else 'Yes [  ]',
                'If yes, do they have a contrary interest?',
                'Yes [X]' if v(d,'contrary')=='Yes' else 'No [X]' if has_partner else '', lw=6, vw=3)

    if has_partner:
        two_col_row(doc, 'Spouse/Partner forename:', v(d,'pfname'), 'Spouse/Partner surname:', v(d,'plname'))
        two_col_row(doc, 'DOB:', v(d,'pdob'), 'NI number:', v(d,'pni'))

    doc.add_paragraph()
    # Dependants
    dep_table = doc.add_table(rows=4, cols=2)
    dep_table.style = 'Table Grid'
    deps = [
        ('How many dependants, currently living with your client, excluding any spouse/partner?', v(d,'dep1')),
        ('How many dependants, not currently living with your client?', v(d,'dep2')),
        ('How many dependants, currently living with your client, does their partner have?', v(d,'dep3') or '0'),
        ('How many dependants, not currently living with your client, does their partner have?', v(d,'dep4') or '0'),
    ]
    for i, (q, a) in enumerate(deps):
        dep_table.rows[i].cells[0].paragraphs[0].add_run(q).font.size = Pt(9)
        dep_table.rows[i].cells[1].paragraphs[0].add_run(a).font.size = Pt(9)

    page_break(doc)

    # ── PAGE 5 ──────────────────────────────────────────────────
    p_bank = doc.add_paragraph()
    p_bank.add_run('Please give details of your client and/or their partner\'s bank, building society and post office accounts:').font.size = Pt(9)

    bank_table = doc.add_table(rows=5, cols=4)
    bank_table.style = 'Table Grid'
    shade_row(bank_table.rows[0])
    for cell, hdr in zip(bank_table.rows[0].cells,
                          ['Bank/building society', 'Account number (last four digits only)', 'Type of account', 'Current balance']):
        r = cell.paragraphs[0].add_run(hdr)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    banks = [
        (v(d,'ba1n'), v(d,'ba1num'), v(d,'ba1t'), v(d,'ba1b')),
        (v(d,'ba2n'), v(d,'ba2num'), v(d,'ba2t'), v(d,'ba2b')),
        (v(d,'ba3n'), v(d,'ba3num'), v(d,'ba3t'), v(d,'ba3b')),
        ('', '', '', ''),
    ]
    for i, (nm, num, typ, bal) in enumerate(banks):
        for j, val in enumerate([nm, num, typ, ('£'+bal if bal else '')]):
            bank_table.rows[i+1].cells[j].paragraphs[0].add_run(val).font.size = Pt(9)

    doc.add_paragraph()
    section_hdr(doc, 'D.  Capital and any other assets (needed for your client & spouse/partner)')

    cap_table = doc.add_table(rows=7, cols=3)
    cap_table.style = 'Table Grid'
    shade_row(cap_table.rows[0])
    for cell, hdr in zip(cap_table.rows[0].cells, ['', 'Client', 'Partner']):
        r = cell.paragraphs[0].add_run(hdr)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    cap_rows = [
        ('Cash (coins, banknotes, cheques)', v(d,'cap_cc'), v(d,'cap_cp')),
        ('Money in banks or building societies', v(d,'cap_bank_c') if d.get('cap_bank_c') else '', ''),
        ('Value of property owned (other than your main house)', v(d,'cap_pc'), v(d,'cap_pp')),
        ('Address(es) of other property: ' + v(d,'prop_addr'), '', ''),
        ('Outstanding value of mortgage/loan secured over other property/land', v(d,'cap_mc') or '', v(d,'cap_mp') or ''),
        ('Investments (shares, bonds, ISAs etc)', v(d,'cap_ic'), v(d,'cap_ip')),
    ]
    for i, (lbl, cv, pv) in enumerate(cap_rows):
        cap_table.rows[i+1].cells[0].paragraphs[0].add_run(lbl).font.size = Pt(8.5)
        cap_table.rows[i+1].cells[1].paragraphs[0].add_run('£'+cv if cv else '£').font.size = Pt(9)
        cap_table.rows[i+1].cells[2].paragraphs[0].add_run('£'+pv if pv else '£').font.size = Pt(9)

    page_break(doc)

    # ── PAGE 6 ──────────────────────────────────────────────────
    section_hdr(doc, 'E.  Income details (needed for client & spouse/partner) - please specify weekly amounts')

    pass_c = v(d,'pass_c')
    pass_p = v(d,'pass_p')
    p_inc = doc.add_paragraph()
    p_inc.add_run('Passport benefits (verifiable by SLAB) - please tick if applicable').font.size = Pt(9)

    pb_table = doc.add_table(rows=5, cols=3)
    pb_table.style = 'Table Grid'
    shade_row(pb_table.rows[0])
    for cell, hdr in zip(pb_table.rows[0].cells, ['Benefit', 'Client', 'Partner']):
        r = cell.paragraphs[0].add_run(hdr)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    pb_items = [
        ('Income Support', 'Income Support'),
        ('Income-based Jobseeker\'s Allowance', 'Income-based JSA'),
        ('Income-related Employment and Support Allowance', 'Income-related ESA'),
        ('Universal Credit', 'Universal Credit'),
    ]
    for i, (label, key) in enumerate(pb_items):
        pb_table.rows[i+1].cells[0].paragraphs[0].add_run(label).font.size = Pt(9)
        pb_table.rows[i+1].cells[1].paragraphs[0].add_run('[X]' if key in pass_c else '[   ]').font.size = Pt(9)
        pb_table.rows[i+1].cells[2].paragraphs[0].add_run('[X]' if key in pass_p else '[   ]').font.size = Pt(9)

    nonpass = v(d,'nonpass')
    p_np = doc.add_paragraph()
    p_np.add_run('Non-passport benefits and other benefits (not verifiable by SLAB)').font.size = Pt(9)
    p_np.runs[0].bold = True

    npb_table = doc.add_table(rows=9, cols=3)
    npb_table.style = 'Table Grid'
    shade_row(npb_table.rows[0])
    for cell, hdr in zip(npb_table.rows[0].cells, ['Benefit', 'Client', 'Partner']):
        r = cell.paragraphs[0].add_run(hdr)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    npb_items = [
        'Contribution-based Jobseeker\'s Allowance',
        'Contribution-based Employment and Support Allowance',
        'Incapacity (not included in income calculation)',
        'Disability Living Allowance (not included in income calculation)',
        'Personal Independence Payment (not included in income calculation)',
        'Child Tax Credit',
        'Child Benefit',
        'Working Tax Credit',
    ]
    for i, label in enumerate(npb_items):
        npb_table.rows[i+1].cells[0].paragraphs[0].add_run(label).font.size = Pt(9)
        checked = any(x.lower() in nonpass.lower() for x in label.split()[:2])
        npb_table.rows[i+1].cells[1].paragraphs[0].add_run('[X]' if checked else '[   ]').font.size = Pt(9)
        npb_table.rows[i+1].cells[2].paragraphs[0].add_run('[   ]').font.size = Pt(9)

    doc.add_paragraph()
    section_hdr(doc, 'F.  Earnings (client & spouse/partner where appropriate) - please specify weekly amounts')

    earn_table = doc.add_table(rows=3, cols=3)
    earn_table.style = 'Table Grid'
    shade_row(earn_table.rows[0])
    for cell, hdr in zip(earn_table.rows[0].cells, ['', 'Client', 'Partner']):
        r = cell.paragraphs[0].add_run(hdr)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    earn_table.rows[1].cells[0].paragraphs[0].add_run('Pay or sick pay (net)').font.size = Pt(9)
    earn_table.rows[1].cells[1].paragraphs[0].add_run('£'+v(d,'earn_pay') if v(d,'earn_pay') else '£').font.size = Pt(9)
    earn_table.rows[1].cells[2].paragraphs[0].add_run('£').font.size = Pt(9)
    earn_table.rows[2].cells[0].paragraphs[0].add_run('Self-employed weekly drawings').font.size = Pt(9)
    earn_table.rows[2].cells[1].paragraphs[0].add_run('£'+v(d,'earn_se') if v(d,'earn_se') else '£').font.size = Pt(9)
    earn_table.rows[2].cells[2].paragraphs[0].add_run('£').font.size = Pt(9)

    if v(d,'employer_c') or v(d,'employer_p'):
        two_col_row(doc, 'Client employer:', v(d,'employer_c'), 'Partner employer:', v(d,'employer_p'))
    if v(d,'no_income'):
        label_value_row(doc, 'If no income, how supported financially:', v(d,'no_income'))

    # Other income
    if v(d,'oi1_src'):
        oi_table = doc.add_table(rows=3, cols=3)
        oi_table.style = 'Table Grid'
        shade_row(oi_table.rows[0])
        for cell, hdr in zip(oi_table.rows[0].cells, ['Other income (please specify)', 'How much?', 'How often?']):
            r = cell.paragraphs[0].add_run(hdr)
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for row_i, (src, amt, freq) in enumerate([
            (v(d,'oi1_src'), v(d,'oi1_amt'), v(d,'oi1_freq')),
            (v(d,'oi2_src'), v(d,'oi2_amt'), v(d,'oi2_freq')),
        ]):
            oi_table.rows[row_i+1].cells[0].paragraphs[0].add_run(src).font.size = Pt(9)
            oi_table.rows[row_i+1].cells[1].paragraphs[0].add_run('£'+amt if amt else '').font.size = Pt(9)
            oi_table.rows[row_i+1].cells[2].paragraphs[0].add_run(freq).font.size = Pt(9)

    label_value_row(doc, 'What documentary evidence was shown to you of your client\'s income and capital?', v(d,'evidence'))

    page_break(doc)

    # ── PAGE 7: Declaration ──────────────────────────────────────
    section_hdr(doc, 'Applicant\'s Declaration and Authority')
    p_d = doc.add_paragraph()
    p_d.add_run('Please read each of the following statements carefully and ask your solicitor to explain anything you do not understand before signing this declaration.').bold = True

    decl_items(doc, [
        'This is a true statement of my personal and financial circumstances.',
        'I understand that if I give false information to the Scottish Legal Aid Board ("SLAB"), I may be prosecuted.',
        'I understand that SLAB can make any enquiries and get any information it needs to deal with this application.',
        'I agree to SLAB obtaining and/or checking information with others such as my employer, banks, credit reference agencies, the Department for Work and Pensions and HM Revenue and Customs and I authorise those people/organisations to provide the information they are asked for.',
        'I agree to the disclosure of the application, associated documentation and my case file held by my solicitor, to SLAB for audit and/or quality assurance.',
        'SLAB may use the information I or my solicitor have provided on this form, or otherwise provide, for the prevention and detection of fraud.',
        'SLAB may share this information with other bodies responsible for auditing or administering public funds for these purposes. I consent to SLAB disclosing my personal data to other organisations.',
        'I agree that all of the above consents and agreements will be effective for a period of not less than five years from the date of signature and any further reasonable period thereafter as SLAB considers appropriate for their requirements.',
    ])

    sig_block(doc, 'Signature of applicant/representative:', client_sig_path, v(d,'sigdate'), width_inches=3.0)

    section_hdr(doc, 'Solicitor\'s Declaration')
    decl_items(doc, [
        'I consent to the disclosure of the application, associated documentation & client case file for quality assurance including audit & peer review, at any stage.',
        'I accept responsibility for any act or omission in relation to the completion & submission of the application on Legal Aid Online ("LAOL") by me or on my behalf & confirm that all information contained within this declaration will be submitted fully & accurately in the online application.',
        'I have satisfied myself that my client qualifies financially for advice & assistance. I have seen either documentary evidence to support my decision or made necessary enquiries with my client to be satisfied they are financially eligible & I have complied with paragraphs 2A(2)(3) of Schedule 2 to the Advice & Assistance (Scotland) Regulations 1996 in relation to verification & otherwise.',
        'I will retain this signed, completed document in paper form or electronically (see LAOL Terms & Conditions for more details) & will send it to SLAB upon request.',
    ])
    sig_block(doc, 'Signature of solicitor:', SOL_SIG_PATH, v(d,'sigdate'), width_inches=3.0)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def make_civ(d, client_sig_path, partner_sig_path=None):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(9)

    # ── PAGE 1 ──────────────────────────────────────────────────
    hdr_table = doc.add_table(rows=1, cols=3)
    hdr_table.style = 'Table Grid'
    hdr_table.rows[0].cells[0].paragraphs[0].add_run('Civil Legal Aid\nLegal Aid Online Declaration').font.size = Pt(11)
    sm_cell = hdr_table.rows[0].cells[1]
    sm_cell.paragraphs[0].add_run('Insert subject matter:').font.size = Pt(8)
    sm_cell.add_paragraph().add_run(v(d,'subject')).font.size = Pt(9)
    hdr_table.rows[0].cells[2].paragraphs[0].add_run('June 2022\nCIV/SOL').font.size = Pt(8)

    p_intro = doc.add_paragraph()
    p_intro.add_run('This Declaration must be completed and signed by the solicitor, the applicant and where applicable, the partner. The information contained in this form will be entered into Legal Aid Online.').font.size = Pt(8.5)

    doc.add_paragraph()
    section_hdr(doc, 'A.  Applicant details')

    two_col_row(doc, 'Forename:', v(d,'fname'), 'Surname:', v(d,'lname'))
    label_value_row(doc, 'Date of birth (dd/mm/yyyy):', v(d,'dob'), 5, 11.5)
    two_col_row(doc, 'Contact telephone number:', v(d,'tel'), 'Contact email address:', v(d,'email'))
    two_col_row(doc, 'Contact by email?', v(d,'cemail'), 'National Insurance number:', v(d,'ni'))
    label_value_row(doc, 'If no NI number, reason:', v(d,'ni_reason'))
    two_col_row(doc, 'Home address:', v(d,'haddr')+'\n'+v(d,'hpc'), 'Correspondence address:', v(d,'caddr')+'\n'+v(d,'cpc'), lw=4, vw=5)

    comm_val = v(d,'comm')
    p_comm = doc.add_paragraph()
    p_comm.add_run('Does your client have any communication support needs?').bold = True
    comm_items = [
        ('none/no support needed', 'None' in comm_val or 'none' in comm_val.lower()),
        ('audio', 'Audio' in comm_val),
        ('spoken language translation or interpreter', 'Spoken' in comm_val),
        ('large print', 'Large print' in comm_val),
        ('British Sign Language (BSL)', 'BSL' in comm_val or 'British Sign' in comm_val),
        ('braille', 'Braille' in comm_val or 'braille' in comm_val),
        ('other - provide details below', 'Other' in comm_val),
    ]
    checkbox_row(doc, comm_items, cols=2)

    page_break(doc)

    # ── PAGE 2 ──────────────────────────────────────────────────
    section_hdr(doc, 'Applicant/client equality information', grey=True)

    sex = v(d,'sex')
    p_sex = doc.add_paragraph()
    p_sex.add_run('What is your client\'s sex?').bold = True
    checkbox_row(doc, [
        ('Female', sex == 'Female'),
        ('Male', sex == 'Male'),
        ('Prefer not to say', sex == 'Prefer not to say'),
    ])

    dis = v(d,'disabilities')
    p_dis = doc.add_paragraph()
    p_dis.add_run('Disabilities').bold = True
    dis_items = [
        ('Deafness or partial hearing loss', 'Deafness' in dis),
        ('Blindness or partial sight loss', 'Blindness' in dis),
        ('Full or partial loss of voice or difficulty speaking', 'voice' in dis.lower()),
        ('Learning disability', 'Learning disability' in dis),
        ('Learning difficulty', 'Learning difficulty' in dis),
        ('Development disorder', 'Development' in dis),
        ('Physical disability', 'Physical' in dis),
        ('Mental health condition', 'Mental' in dis),
        ('Long-term illness, disease or condition', 'Long-term' in dis),
        ('Other condition: please write in below', 'Other condition' in dis),
        ('No known conditions', 'No known' in dis),
        ('Prefer not to say', dis == 'Prefer not to say'),
    ]
    checkbox_row(doc, dis_items)

    natid = v(d,'natid')
    p_nat = doc.add_paragraph()
    p_nat.add_run('National identity').bold = True
    checkbox_row(doc, [
        ('Scottish', 'Scottish' in natid),
        ('English', 'English' in natid),
        ('Northern Irish', 'Northern Irish' in natid),
        ('Welsh', 'Welsh' in natid),
        ('British', 'British' in natid),
        ('Prefer not to say', natid == 'Prefer not to say'),
    ], cols=2)

    page_break(doc)

    # ── PAGE 3 ──────────────────────────────────────────────────
    ethnic = v(d,'ethnic')
    p_eth = doc.add_paragraph()
    p_eth.add_run('Ethnic group').bold = True
    eth_items = [
        ('White Scottish', 'White Scottish' in ethnic),
        ('White Other British', 'White Other British' in ethnic),
        ('White Irish', 'White Irish' in ethnic),
        ('White Polish', 'White Polish' in ethnic),
        ('White Gypsy/Traveller', 'Gypsy' in ethnic),
        ('White Roma', 'Roma' in ethnic),
        ('White Showman/Showwoman', 'Showman' in ethnic),
        ('Pakistani, Scottish Pakistani or British Pakistani', 'Pakistani' in ethnic),
        ('Indian, Scottish Indian or British Indian', 'Indian' in ethnic),
        ('Bangladeshi', 'Bangladeshi' in ethnic),
        ('Chinese, Scottish Chinese or British Chinese', 'Chinese' in ethnic),
        ('African, Scottish African or British African', 'African' in ethnic),
        ('Caribbean or Black', 'Caribbean' in ethnic),
        ('Mixed or multiple ethnic groups', 'Mixed' in ethnic),
        ('Arab, Scottish Arab or British Arab', 'Arab' in ethnic),
        ('Prefer not to say', ethnic == 'Prefer not to say'),
    ]
    checkbox_row(doc, eth_items, cols=2)

    care = v(d,'care')
    p_care = doc.add_paragraph()
    p_care.add_run('Care Experience').bold = True
    checkbox_row(doc, [
        ("Currently 'looked after' by Local Authority", 'Currently' in care),
        ("Have previously been 'looked after' by Local Authority", 'previously' in care.lower()),
        ("Never been 'looked after' by Local Authority", 'Never' in care),
        ('Prefer not to say', care == 'Prefer not to say'),
    ])

    page_break(doc)

    # ── PAGE 4 ──────────────────────────────────────────────────
    section_hdr(doc, 'B.  Other rights and resources')
    oa = v(d,'otherassist') == 'Yes'
    two_col_row(doc, 'Does your client have access to other assistance providing help with legal costs?',
                'Yes [X]' if oa else 'Yes [  ]', '', 'No [X]' if not oa else 'No [  ]', lw=11, vw=1.5)
    if oa and v(d,'oa_why'):
        label_value_row(doc, 'If this assistance cannot be used, tell us why:', v(d,'oa_why'))

    section_hdr(doc, 'C.  Financial details  Passported benefits only (all other applicants must complete Form 2)')
    has_partner = v(d,'partner') == 'Yes'
    two_col_row(doc, 'Does your client live with a spouse/partner?',
                'Yes [X]' if has_partner else 'No [X]',
                'If yes, do they have a contrary interest?',
                'Yes [X]' if v(d,'contrary')=='Yes' else 'No [X]', lw=6, vw=3)

    pass_c = v(d,'pass_c')
    pass_p = v(d,'pass_p')
    pb_table = doc.add_table(rows=5, cols=3)
    pb_table.style = 'Table Grid'
    shade_row(pb_table.rows[0])
    for cell, hdr in zip(pb_table.rows[0].cells, ['Does your client or their partner receive any of these benefits?', 'Client', 'Partner']):
        r = cell.paragraphs[0].add_run(hdr)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, (label, key) in enumerate([
        ('Income Support', 'Income Support'),
        ('Income-based Jobseeker\'s Allowance', 'Income-based JSA'),
        ('Income-related Employment and Support Allowance', 'Income-related ESA'),
        ('Universal Credit', 'Universal Credit'),
    ]):
        pb_table.rows[i+1].cells[0].paragraphs[0].add_run(label).font.size = Pt(9)
        pb_table.rows[i+1].cells[1].paragraphs[0].add_run('[X]' if key in pass_c else '[   ]').font.size = Pt(9)
        pb_table.rows[i+1].cells[2].paragraphs[0].add_run('[X]' if key in pass_p else '[   ]').font.size = Pt(9)

    doc.add_paragraph()
    p_oi_hdr = doc.add_paragraph()
    p_oi_hdr.add_run('Does your client or their partner have income from another source?').bold = True

    oi_table = doc.add_table(rows=4, cols=3)
    oi_table.style = 'Table Grid'
    shade_row(oi_table.rows[0])
    for cell, hdr in zip(oi_table.rows[0].cells, ['Other income', 'How much?', 'How often?']):
        r = cell.paragraphs[0].add_run(hdr)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    oi_data = [
        (v(d,'oi1_src'), v(d,'oi1_amt'), v(d,'oi1_freq')),
        (v(d,'oi2_src'), v(d,'oi2_amt'), v(d,'oi2_freq')),
        ('', '', ''),
    ]
    for i, (src, amt, freq) in enumerate(oi_data):
        oi_table.rows[i+1].cells[0].paragraphs[0].add_run(src).font.size = Pt(9)
        oi_table.rows[i+1].cells[1].paragraphs[0].add_run('£'+amt if amt else '£').font.size = Pt(9)
        oi_table.rows[i+1].cells[2].paragraphs[0].add_run(freq).font.size = Pt(9)

    page_break(doc)

    # ── PAGE 5: Applicant Declaration ───────────────────────────
    section_hdr(doc, 'Applicant\'s declaration and Authority')
    p_d = doc.add_paragraph()
    p_d.add_run('Please read each of the following statements carefully and ask your solicitor to explain anything you do not understand before signing this declaration').bold = True

    decl_items(doc, [
        'This is a true statement of my personal and financial circumstances.',
        'I understand that if I give false information to the Scottish Legal Aid Board ("SLAB"), I may be prosecuted.',
        'I understand that SLAB can make any enquiries and get any information it needs to deal with this application.',
        'I agree to SLAB obtaining and/or checking information with others such as my employer, banks, credit reference agencies, the Department for Work and Pensions and HM Revenue and Customs and I authorise those people/organisations to provide the information they are asked for.',
        'I understand that I must tell you immediately if there are any changes in my or my partner\'s financial circumstances including a change in benefits. Failure to notify changes may lead to my legal aid being revoked and I may be prosecuted. I may also be liable to pay the costs of my case.',
        'If my solicitor does special urgency work for me I know that SLAB may need me to pay a contribution towards that work. I agree to pay any contribution assessed by SLAB on the information I have provided.',
        'I agree to the disclosure of the application, associated documentation and my case file held by my solicitor, to SLAB for audit and/or quality assurance.',
        'SLAB may use the information I or my solicitor have provided on this form, or otherwise provide, for the prevention and detection of fraud.',
        'SLAB may share this information with other bodies responsible for auditing or administering public funds for these purposes. I consent to SLAB disclosing my personal data to other organisations.',
        'I agree that all of the above consents and agreements will be effective for a period of not less than five years from the date of signature and any further reasonable period thereafter as SLAB considers appropriate for their requirements.',
    ])

    sig_block(doc, 'Signature of applicant/representative:', client_sig_path, v(d,'sigdate'), width_inches=3.0)

    page_break(doc)

    # ── PAGE 6: Solicitor + Partner Declaration ──────────────────
    section_hdr(doc, 'Solicitor\'s Declaration')
    decl_items(doc, [
        'I consent to the disclosure of the application, associated documentation and client case file for quality assurance including audit and peer review, at any stage.',
        'I accept responsibility for any act or omission in relation to the completion and submission of the application on Legal Aid Online ("LAOL") by me or on my behalf and confirm that all information contained within this declaration will be submitted fully and accurately in the online application.',
        'I will retain this signed, completed document in paper form or electronically (see LAOL Terms and Conditions for more details) and will send it to SLAB upon request.',
    ])
    sig_block(doc, 'Signature of solicitor:', SOL_SIG_PATH, v(d,'sigdate'), width_inches=3.0)

    section_hdr(doc, 'Partner\'s declaration (complete where an online Form 1 will be completed)')
    if has_partner:
        two_col_row(doc, 'Spouse/partner name:', v(d,'pfname')+' '+v(d,'plname'), 'Date of birth:', v(d,'pdob'))
        label_value_row(doc, 'NI Number:', v(d,'pni'))
    else:
        label_value_row(doc, 'Spouse/partner name:', '')
        label_value_row(doc, 'Date of birth (dd/mm/yyyy):', '')
        label_value_row(doc, 'NI Number:', '')

    decl_items(doc, [
        'I have seen the financial information in this application.',
        'This is a true statement of my personal and financial circumstances.',
        'I understand that if I give false information to the Scottish Legal Aid Board (SLAB), I may be prosecuted.',
        'I agree to SLAB checking these facts with others such as banks, credit reference agencies, the Department of Work and Pensions and HM Revenue and Customs and I authorise those people/organisations to provide the information they are asked for.',
    ])

    p_ps = doc.add_paragraph()
    two_col_row(doc, 'Signature of applicant\'s partner:', '', 'Date:', v(d,'sigdate'))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def send_email(d, aa_buf, civ_buf):
    lname = v(d,'lname')
    fname = v(d,'fname')
    ref = v(d,'ref')
    body = ('New SLAB declaration submitted online.\n\n'
            'Client: ' + fname + ' ' + lname + '\n'
            'DOB: ' + v(d,'dob') + '\n'
            'NI: ' + (v(d,'ni') or 'Not provided') + '\n'
            'Tel: ' + v(d,'tel') + '\n'
            'Email: ' + v(d,'email') + '\n'
            'Address: ' + v(d,'haddr') + ', ' + v(d,'hpc') + '\n'
            'Subject Matter: ' + v(d,'subject') + '\n'
            'Passported Benefits: ' + v(d,'pass_c') + '\n'
            'Evidence Provided: ' + v(d,'evidence') + '\n'
            'Date Signed: ' + v(d,'sigdate') + '\n'
            'Reference: ' + ref + '\n\n'
            'Both completed declaration documents are attached as PDFs.\n'
            'Please countersign and submit to SLAB via Legal Aid Online.')

    msg = MIMEMultipart()
    msg['From'] = GMAIL_ADDRESS
    msg['To'] = FIRM_EMAIL
    msg['Subject'] = 'New Legal Aid Declaration - ' + fname + ' ' + lname + ' - ' + ref
    msg.attach(MIMEText(body, 'plain'))

    for buf, filename in [
        (aa_buf, 'AA_LAO_CIV_' + lname + '_' + fname + '.docx'),
        (civ_buf, 'CIV_SOL_' + lname + '_' + fname + '.docx'),
    ]:
        part = MIMEBase('application', 'octet-stream')
        part.set_payload(buf.read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', 'attachment; filename=' + filename)
        msg.attach(part)

    server = smtplib.SMTP('smtp.gmail.com', 587)
    server.starttls()
    server.login(GMAIL_ADDRESS, GMAIL_APP_PASSWORD)
    server.send_message(msg)
    server.quit()


@app.route('/submit', methods=['POST'])
def submit():
    try:
        d = request.json
        d['ref'] = 'TFL-' + str(int(datetime.now().timestamp()))[-6:]
        print('Submit:', v(d,'fname'), v(d,'lname'))

        client_sig_path = make_client_sig_image(v(d,'signame'), d.get('sigDataURL',''))
        partner_sig_path = make_client_sig_image(v(d,'pfname')+' '+v(d,'plname'), d.get('partnerSigDataURL','')) if v(d,'partner')=='Yes' else None

        aa_buf = make_aa(d, client_sig_path)
        print('AA built')
        civ_buf = make_civ(d, client_sig_path, partner_sig_path)
        print('CIV built')

        send_email(d, aa_buf, civ_buf)
        print('Email sent')

        os.unlink(client_sig_path)
        if partner_sig_path:
            os.unlink(partner_sig_path)
        return jsonify({'ok': True, 'ref': d['ref']})
    except Exception as e:
        import traceback
        print('ERROR:', str(e))
        traceback.print_exc()
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/')
def home():
    return 'SLAB Forms Server - OK'


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 5000)))
